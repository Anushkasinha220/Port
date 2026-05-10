from __future__ import annotations

import json
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "report_assets"
OUT = ROOT / "Port_Optimus_AAT_Report.docx"
ACCENT = "0F766E"
INK = "0F172A"
MUTED = "64748B"
BORDER = "CBD5E1"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/consolab.ttf" if bold else "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def safe_text(text: str) -> str:
    return text.replace("⚓", "P").replace("🚢", "SHIP").replace("📦", "BOX").replace("🚚", "TRUCK")


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill: str = "E2E8F0") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if row_i == 0:
                add_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_caption(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(label)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(100, 116, 139)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def draw_arrow(draw, start, end, color="#334155", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
    else:
        pts = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(pts, fill=color)


def box(draw, xy, title, subtitle="", fill="#ffffff", outline="#CBD5E1", title_color="#0F172A"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    draw.text((x1 + 18, y1 + 16), title, font=font(24, True), fill=title_color)
    if subtitle:
        lines = textwrap.wrap(subtitle, width=28)
        y = y1 + 50
        for line in lines[:3]:
            draw.text((x1 + 18, y), line, font=font(17), fill="#475569")
            y += 23


def create_architecture_diagram():
    img = Image.new("RGB", (1600, 900), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((50, 36), "Port-Optimus System Architecture", font=font(34, True), fill="#0F172A")
    boxes = {
        "env": (70, 150, 420, 320),
        "state": (530, 150, 900, 320),
        "agent": (1010, 150, 1450, 320),
        "actions": (530, 430, 900, 620),
        "mlops": (1010, 430, 1450, 620),
        "dash": (70, 430, 420, 620),
        "sdg": (410, 705, 1110, 830),
    }
    box(d, boxes["env"], "sim/port_env.py", "Congestion, arrivals, prices, battery, reward")
    box(d, boxes["state"], "State Vector", "Queue, price, battery, trucks, hour")
    box(d, boxes["agent"], "src/engine.py", "Granular Q-learning with epsilon decay; optional DQN")
    box(d, boxes["actions"], "Action Space", "0 Rapid Discharge, 1 Buffered Move, 2 Idle / Charge")
    box(d, boxes["mlops"], "src/mlops.py", "Run ID, metadata.json, results.csv")
    box(d, boxes["dash"], "app.py", "Streamlit control center with Plotly analytics")
    box(d, boxes["sdg"], "SDG 9, 11, 12 Impact", "Efficient infrastructure, sustainable cities, responsible energy use", fill="#ECFEFF", outline="#14B8A6")
    draw_arrow(d, (420, 235), (530, 235))
    draw_arrow(d, (900, 235), (1010, 235))
    draw_arrow(d, (1230, 320), (760, 430))
    draw_arrow(d, (530, 525), (420, 525))
    draw_arrow(d, (1230, 320), (1230, 430))
    draw_arrow(d, (240, 620), (590, 705), color="#0F766E")
    draw_arrow(d, (1230, 620), (930, 705), color="#0F766E")
    img.save(ASSET_DIR / "architecture.png")


def create_workflow_diagram():
    img = Image.new("RGB", (1600, 650), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((50, 32), "Training and Simulation Workflow", font=font(34, True), fill="#0F172A")
    labels = [
        ("1. Reset Environment", "Initial queue, battery, trucks, hour"),
        ("2. Observe State", "[Queue, price, battery, trucks, hour]"),
        ("3. Select Action", "Epsilon-greedy Q policy"),
        ("4. Step Port", "Process, charge, update congestion"),
        ("5. Compute Reward", "Throughput - cost - wait penalty"),
        ("6. Log Run", "Manifest + results CSV"),
    ]
    x = 60
    for i, (title, sub) in enumerate(labels):
        box(d, (x, 170, x + 230, 365), title, sub, fill="#F8FAFC")
        if i < len(labels) - 1:
            draw_arrow(d, (x + 230, 267), (x + 290, 267), color="#0F766E")
        x += 250
    d.rounded_rectangle((190, 455, 1410, 555), radius=12, fill="#F0FDFA", outline="#14B8A6", width=2)
    d.text((230, 482), "Feedback loop: Q(s,a) is updated after each action, epsilon decays after each episode, and policy quality is measured against a fixed human baseline.", font=font(22, True), fill="#115E59")
    img.save(ASSET_DIR / "workflow.png")


def create_dashboard_layout():
    img = Image.new("RGB", (1600, 900), "#EEF2F6")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((40, 40, 300, 860), radius=14, fill="#0B1220")
    d.text((70, 80), "P  Port-Optimus", font=font(27, True), fill="#FFFFFF")
    for i, label in enumerate(["Trigger Storm", "Peak Season", "Reset Tests", "Log runs"]):
        y = 160 + i * 70
        d.rounded_rectangle((70, y, 270, y + 45), radius=8, fill="#111827", outline="#334155")
        d.text((90, y + 12), label, font=font(17, True), fill="#E2E8F0")
    d.text((340, 52), "Port-Optimus Logistics & Energy Grid", font=font(36, True), fill="#0F172A")
    d.text((340, 96), "Autonomous port command center", font=font(19), fill="#64748B")
    for i, label in enumerate(["Scenario", "Battery", "Throughput", "Energy", "Safety"]):
        x = 340 + i * 240
        d.rounded_rectangle((x, 150, x + 215, 255), radius=10, fill="#FFFFFF", outline="#D8DEE8")
        d.text((x + 18, 170), label.upper(), font=font(15, True), fill="#64748B")
        d.text((x + 18, 198), ["Normal", "72%", "2.88", "0.35", "92"][i], font=font(30, True), fill="#0F172A")
    d.rounded_rectangle((340, 285, 1025, 585), radius=10, fill="#FFFFFF", outline="#D8DEE8")
    d.text((365, 310), "Live Port Digital Twin", font=font(24, True), fill="#0F172A")
    for y, label, icons in [(370, "BERTH QUEUE", "SHIP  SHIP  SHIP"), (445, "BUFFER GRID", "BOX  BOX  BOX  BOX  BOX"), (520, "E-TRUCK LANE", "TRUCK  TRUCK  TRUCK")]:
        d.text((365, y), label, font=font(16, True), fill="#334155")
        d.text((560, y - 4), icons, font=font(22, True), fill="#0F766E")
    d.rounded_rectangle((1050, 285, 1515, 585), radius=10, fill="#0F172A")
    d.text((1080, 315), "Policy Decision Engine", font=font(24, True), fill="#FFFFFF")
    d.text((1080, 365), "Recommended: Buffered Move", font=font(27, True), fill="#F8FAFC")
    for i, label in enumerate(["Energy Used", "Grid Cost", "Wait Penalty", "Safety"]):
        x = 1080 + (i % 2) * 200
        y = 425 + (i // 2) * 70
        d.rounded_rectangle((x, y, x + 170, y + 50), radius=8, fill="#182235", outline="#263244")
        d.text((x + 12, y + 8), label, font=font(13, True), fill="#94A3B8")
    for x, title in [(340, "RL Radar"), (750, "Energy Heatmap"), (1160, "Policy Trace")]:
        d.rounded_rectangle((x, 625, x + 355, 835), radius=10, fill="#FFFFFF", outline="#D8DEE8")
        d.text((x + 18, 650), title, font=font(23, True), fill="#0F172A")
        d.line((x + 24, 710, x + 330, 710), fill="#CBD5E1", width=3)
        d.line((x + 24, 755, x + 300, 755), fill="#14B8A6", width=6)
    img.save(ASSET_DIR / "dashboard_layout.png")


def create_metric_chart():
    metrics = json.loads((ROOT / "report_metrics.json").read_text(encoding="utf-8"))
    img = Image.new("RGB", (1600, 850), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((55, 40), "Scenario Results: RL Agent vs Fixed Baseline", font=font(34, True), fill="#0F172A")
    y0 = 150
    max_gain = max(item["throughput_gain_pct"] for item in metrics)
    for i, item in enumerate(metrics):
        y = y0 + i * 150
        d.text((70, y), item["scenario"], font=font(24, True), fill="#0F172A")
        d.text((70, y + 36), f"Energy intensity reduction: {item['energy_intensity_reduction_pct']}%", font=font(18), fill="#475569")
        d.text((70, y + 65), f"Wait penalty RL {item['RL_Average_Wait_Penalty']} vs baseline {item['Baseline_Average_Wait_Penalty']}", font=font(18), fill="#475569")
        bar_w = int(800 * item["throughput_gain_pct"] / max_gain)
        d.rounded_rectangle((570, y + 15, 1370, y + 55), radius=9, fill="#E2E8F0")
        d.rounded_rectangle((570, y + 15, 570 + bar_w, y + 55), radius=9, fill="#0F766E")
        d.text((1395, y + 18), f"+{item['throughput_gain_pct']}%", font=font(21, True), fill="#0F766E")
        d.text((570, y + 66), f"Throughput: RL {item['RL_Throughput_Rate']} /hr  |  Baseline {item['Baseline_Throughput_Rate']} /hr", font=font(18), fill="#334155")
    img.save(ASSET_DIR / "metrics_chart.png")


def create_vs_tree():
    lines = [
        "Port_Optimus_VS/",
        "├── app.py                         # Streamlit dashboard",
        "├── sim/",
        "│   └── port_env.py                # RL environment",
        "├── src/",
        "│   ├── engine.py                  # Q-learning + optional DQN",
        "│   ├── mlops.py                   # manifest + CSV logging",
        "│   └── train_port_optimus.py      # training entrypoint",
        "├── models/port_optimus_q_agent.pkl",
        "├── experiments/metadata.json",
        "├── experiments/results.csv",
        "└── .vscode/launch.json            # VS Code run profile",
    ]
    create_code_image("VS Code - Project Explorer", "\n".join(lines), ASSET_DIR / "vs_project_tree.png", show_lines=False)


def extract_snippet(path: str, start: int, end: int) -> str:
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    selected = []
    for idx in range(start, min(end, len(lines)) + 1):
        selected.append(f"{idx:>3}  {safe_text(lines[idx - 1])}")
    return "\n".join(selected)


def create_code_image(title: str, code: str, out: Path, show_lines: bool = True) -> None:
    code = safe_text(code)
    lines = code.splitlines()
    width = 1500
    line_h = 27
    height = 92 + max(8, len(lines)) * line_h + 36
    img = Image.new("RGB", (width, height), "#0B1220")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((28, 28, width - 28, height - 28), radius=14, fill="#111827", outline="#334155", width=2)
    d.rectangle((28, 28, width - 28, 82), fill="#1E293B")
    for i, c in enumerate(["#EF4444", "#F59E0B", "#22C55E"]):
        d.ellipse((54 + i * 30, 47, 68 + i * 30, 61), fill=c)
    d.text((160, 46), title, font=font(21, True), fill="#E2E8F0")
    y = 105
    for line in lines:
        color = "#CBD5E1"
        stripped = line.strip()
        if stripped.startswith("def ") or stripped.startswith("class ") or "def " in stripped:
            color = "#5EEAD4"
        elif stripped.startswith("#") or "# " in stripped:
            color = "#94A3B8"
        elif "return" in stripped or "if " in stripped or "else" in stripped:
            color = "#FBBF24"
        d.text((62, y), line[:132], font=font(18), fill=color)
        y += line_h
    img.save(out)


def create_snippets():
    create_vs_tree()
    create_code_image(
        "VS Code - sim/port_env.py | Custom Environment",
        extract_snippet("sim/port_env.py", 33, 66),
        ASSET_DIR / "snippet_env.png",
    )
    create_code_image(
        "VS Code - src/engine.py | Epsilon-Greedy Q-Learning",
        extract_snippet("src/engine.py", 35, 75),
        ASSET_DIR / "snippet_engine.png",
    )
    create_code_image(
        "VS Code - src/mlops.py | Manifest + Results Logging",
        extract_snippet("src/mlops.py", 20, 58),
        ASSET_DIR / "snippet_mlops.png",
    )
    create_code_image(
        "VS Code - app.py | Live Simulation Loop",
        extract_snippet("app.py", 511, 540),
        ASSET_DIR / "snippet_app.png",
    )


def make_assets():
    ASSET_DIR.mkdir(exist_ok=True)
    create_architecture_diagram()
    create_workflow_diagram()
    create_dashboard_layout()
    create_metric_chart()
    create_snippets()


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "Port-Optimus AAT Report"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "Autonomous Logistics & Energy Grid"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(21)
    styles["Title"].font.bold = True
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor(15, 23, 42)


def cover_page(doc: Document):
    for text, size, bold in [
        ("B. M. S. COLLEGE OF ENGINEERING", 16, True),
        ("Autonomous Institute, Affiliated to VTU, Belagavi", 11, False),
        ("DEPARTMENT OF MACHINE LEARNING", 13, True),
        ("Program: B.E. in Artificial Intelligence and Machine Learning", 11, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AAT / Lab Report")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 118, 110)
    doc.add_paragraph()
    topic = (
        "Port-Optimus: Autonomous Logistics & Energy Grid - Reinforcement Learning "
        "System for Sustainable Port Container Routing"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Topic\n")
    r.bold = True
    r.font.size = Pt(13)
    r2 = p.add_run(topic)
    r2.font.size = Pt(12)
    r2.italic = True
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    rows = [
        ("Submitted by", "Anushka Sinha"),
        ("USN", "1BM23AI229"),
        ("Semester & Section", "6D"),
        ("Faculty In-charge", "Prof. Yashaswi"),
        ("Project Folder", r"C:\Users\Anush\Downloads\Internals_Basics\Port_Optimus_VS"),
        ("Date", "May 2026"),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    style_table(table, "E0F2FE")
    doc.add_paragraph()
    valuation = doc.add_table(rows=3, cols=2)
    for i, (k, v) in enumerate([("Score", "(to be filled by faculty)"), ("Comments", ""), ("Faculty Signature with date", "")]):
        valuation.cell(i, 0).text = k
        valuation.cell(i, 1).text = v
    style_table(valuation, "F8FAFC")
    doc.add_page_break()


def toc(doc: Document):
    add_heading(doc, "Table of Contents", 1)
    items = [
        "1. Introduction",
        "1.1 Background and Motivation",
        "1.2 Importance of AI, MLOps, and Interactive Tools in RL",
        "1.3 Objective of the Study: Port-Optimus",
        "2. Topic Explored: Reinforcement Learning for Autonomous Ports",
        "2.1 Description of the Port-Optimus Innovation",
        "2.2 Theoretical Background: Markov Decision Process and Q-Learning",
        "2.3 Mathematical Model: State, Action, Reward, and Update Rule",
        "2.4 Relevance to Sustainable Infrastructure",
        "3. Problem Statement",
        "4. Methodology and Training Strategy",
        "5. Implementation Details with VS Code Snippets",
        "6. Results and Analysis",
        "7. SDG Impact Report",
        "8. Challenges and Limitations",
        "9. Conclusion and Future Work",
        "10. References",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.2 if "." in item[:4] else 0)
    doc.add_page_break()


def add_image(doc: Document, filename: str, caption: str, width: float = 6.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSET_DIR / filename), width=Inches(width))
    add_caption(doc, caption)


def metrics_table(doc: Document):
    data = json.loads((ROOT / "report_metrics.json").read_text(encoding="utf-8"))
    table = doc.add_table(rows=1, cols=6)
    headers = ["Scenario", "RL Throughput", "Baseline Throughput", "Energy Intensity Reduction", "RL Wait Penalty", "Safety Gain"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for item in data:
        cells = table.add_row().cells
        cells[0].text = item["scenario"]
        cells[1].text = f"{item['RL_Throughput_Rate']} /hr"
        cells[2].text = f"{item['Baseline_Throughput_Rate']} /hr"
        cells[3].text = f"{item['energy_intensity_reduction_pct']}%"
        cells[4].text = f"{item['RL_Average_Wait_Penalty']} vs {item['Baseline_Average_Wait_Penalty']}"
        cells[5].text = f"{round(item['RL_Average_Safety'] - item['Baseline_Average_Safety'], 2)}"
    style_table(table, "E0F2FE")


def build_doc():
    make_assets()
    doc = Document()
    configure_doc(doc)
    cover_page(doc)
    toc(doc)

    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Background and Motivation", 2)
    add_body(doc, "Modern ports are becoming cyber-physical logistics hubs where ship arrivals, container movement, electric truck availability, buffer-yard capacity, battery storage, and electricity prices must be coordinated in real time. Traditional rule-based dispatching often moves containers quickly during expensive energy windows or allows queues to grow during periods of congestion. Port-Optimus addresses this problem by using reinforcement learning to decide when to rapidly discharge energy for fast processing, when to use buffered eco-moves, and when to idle or charge during low-energy windows.")
    add_heading(doc, "1.2 Importance of AI, MLOps, and Interactive Tools in RL", 2)
    add_body(doc, "The project combines AI decision-making, a custom simulator, MLOps run tracking, and an interactive Streamlit dashboard. This mirrors the workflow of practical industrial AI systems: train a policy, evaluate it against a baseline, log reproducible runs, and expose the behaviour through a control interface that can be stress-tested by operators.")
    add_heading(doc, "1.3 Objective of the Study: Port-Optimus", 2)
    add_bullets(doc, [
        "Develop a custom port environment with congestion, fluctuating electricity prices, battery level, truck availability, and ship arrivals.",
        "Train an RL agent using granular Q-learning with epsilon-greedy exploration and decay.",
        "Implement save and load functions for the learned model.",
        "Create an MLOps manifest system that generates run IDs and logs results into CSV.",
        "Build a professional Streamlit dashboard with live simulation, radar chart, heatmap, stress tests, and policy trace.",
        "Evaluate SDG impact for resilient infrastructure, sustainable cities, and responsible energy use.",
    ])

    add_heading(doc, "2. Topic Explored: Reinforcement Learning for Autonomous Ports", 1)
    add_heading(doc, "2.1 Description of the Port-Optimus Innovation", 2)
    add_body(doc, "Port-Optimus is a digital-twin style RL system for autonomous port logistics and energy management. The innovation is not only the reward function, but the integration of logistics throughput, energy price response, congestion risk, battery preservation, and MLOps tracking into a single demonstrable project.")
    add_image(doc, "architecture.png", "Figure 1: System architecture showing environment, state vector, RL engine, actions, MLOps logging, dashboard, and SDG impact.")
    add_heading(doc, "2.2 Theoretical Background: Markov Decision Process and Q-Learning", 2)
    add_body(doc, "The environment is modelled as a Markov Decision Process. At each hour, the agent observes the port state, chooses one action, receives a reward, and transitions to a new state. Q-learning estimates the long-term value of each action in each discretised state, allowing the agent to improve through repeated simulated episodes.")
    add_heading(doc, "2.3 Mathematical Model: State, Action, Reward, and Update Rule", 2)
    add_body(doc, "State = [Queue_Length, Current_Electricity_Price, Port_Battery_Level, Truck_Availability, Hour_of_Day]. The action space contains Rapid Discharge, Buffered Move, and Idle/Charge. The reward is defined as (Units_Processed x 10) - (Energy_Cost x Penalty) - Wait_Time_Penalty. The Q-value update uses Q(s,a) <- Q(s,a) + alpha [r + gamma max_a' Q(s',a') - Q(s,a)].")
    add_heading(doc, "2.4 Relevance to Sustainable Infrastructure", 2)
    add_body(doc, "The project is aligned with SDG 9, SDG 11, and SDG 12 by improving industrial infrastructure intelligence, reducing logistics congestion in urban port regions, and encouraging responsible energy use through price-aware dispatching.")

    add_heading(doc, "3. Problem Statement", 1)
    add_body(doc, "Develop a reinforcement learning system called Port-Optimus: Autonomous Logistics & Energy Grid. The agent must manage an autonomous port where ships arrive with cargo and decide whether to move containers to electric trucks, place them in high-density buffers, or wait for low-energy windows while managing a limited port battery and fluctuating electricity prices.")
    add_bullets(doc, [
        "Input data: simulated 24-hour port operations with queue length, electricity price, battery level, trucks, and hour.",
        "Expected outcome: higher throughput, lower wait penalty, improved energy efficiency, and robust behaviour under storm and peak-season stress tests.",
        "Baseline: a fixed human-style dispatch rule used for performance comparison.",
    ])

    add_heading(doc, "4. Methodology and Training Strategy", 1)
    add_heading(doc, "4.1 Workflow / Steps Followed", 2)
    add_bullets(doc, [
        "Created a custom simulator in sim/port_env.py.",
        "Defined the state vector, action space, congestion mechanic, reward function, and episode metrics.",
        "Implemented granular Q-learning in src/engine.py with epsilon-greedy exploration and decay.",
        "Trained the agent using src/train_port_optimus.py and saved the learned model in models/.",
        "Logged the run through src/mlops.py into experiments/metadata.json and experiments/results.csv.",
        "Designed the Streamlit dashboard in app.py with stress tests, live animation, Plotly radar chart, heatmap, and trace table.",
        "Evaluated the RL policy against a human-fixed baseline across normal and stressed scenarios.",
    ])
    add_image(doc, "workflow.png", "Figure 2: Training and simulation workflow for the Port-Optimus RL loop.")
    add_heading(doc, "4.2 Model Architecture", 2)
    add_body(doc, "The implemented controller is a granular Q-learning agent. Continuous values such as price and battery level are discretised into bins, producing a compact but expressive state table. The code also contains an optional PyTorch DQN implementation, which can be activated when Torch is installed.")
    add_heading(doc, "4.3 Hyperparameters / Settings Used", 2)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 1).text = "Value"
    for k, v in [
        ("Learning rate", "0.12"),
        ("Discount factor", "0.94"),
        ("Initial epsilon", "1.0"),
        ("Minimum epsilon", "0.05"),
        ("Epsilon decay", "0.992"),
        ("Training episodes", "320 in training script"),
        ("Episode length", "24 simulated hours"),
        ("Actions", "3"),
    ]:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v
    style_table(table, "E0F2FE")
    add_heading(doc, "4.4 Training Stability Strategies", 2)
    add_bullets(doc, [
        "Granular discretisation prevents unstable learning from raw continuous state values.",
        "Epsilon decay gradually shifts the policy from exploration to exploitation.",
        "Reward shaping balances throughput, energy cost, and waiting penalty.",
        "Stress tests evaluate robustness under reduced solar input and doubled ship arrivals.",
    ])

    add_heading(doc, "5. Implementation Details with VS Code Snippets", 1)
    add_image(doc, "vs_project_tree.png", "Figure 3: VS Code project explorer view showing the clean Port_Optimus_VS folder structure.")
    add_image(doc, "snippet_env.png", "Figure 4: VS Code snippet from sim/port_env.py showing the custom environment state and action design.")
    add_image(doc, "snippet_engine.png", "Figure 5: VS Code snippet from src/engine.py showing discretisation and epsilon-greedy action selection.")
    add_image(doc, "snippet_mlops.png", "Figure 6: VS Code snippet from src/mlops.py showing run ID generation and manifest logging.")
    add_image(doc, "snippet_app.png", "Figure 7: VS Code snippet from app.py showing the live Streamlit animation loop.")
    add_image(doc, "dashboard_layout.png", "Figure 8: Labelled dashboard layout implemented for the professional Streamlit control center.")

    add_heading(doc, "6. Results and Analysis", 1)
    add_body(doc, "The RL policy was evaluated over 20 random seeds for each scenario. The results indicate that Port-Optimus improves throughput and reduces waiting penalties across all scenarios. Under stress tests, the advantage becomes more visible because the learned policy adapts to congestion and battery limitations better than the fixed baseline.")
    metrics_table(doc)
    add_image(doc, "metrics_chart.png", "Figure 9: Scenario comparison chart showing throughput gains and energy-intensity reduction.")
    add_heading(doc, "6.1 Dashboard Analytics", 2)
    add_body(doc, "The dashboard contains a radar chart comparing speed, cost, energy, and safety; a 24-hour energy heatmap; a live digital twin; a policy decision panel; and a trace table. These visualisations demonstrate multi-objective optimisation instead of only reporting a single reward value.")

    add_heading(doc, "7. SDG Impact Report", 1)
    add_body(doc, "Port-Optimus supports SDG 9 by demonstrating intelligent, resilient industrial infrastructure. It supports SDG 11 by reducing port congestion that affects surrounding urban mobility. It supports SDG 12 by shifting container movement toward more responsible energy windows and reducing waste from idle queues.")
    add_bullets(doc, [
        "Normal scenario: approximately 12.83% reduction in energy intensity per processed container and 35.93% throughput gain versus baseline.",
        "Storm scenario: approximately 18.63% energy-intensity reduction and 83.89% throughput gain when solar input is reduced.",
        "Peak Season scenario: approximately 19.51% energy-intensity reduction and 40.75% throughput gain under doubled ship arrivals.",
        "Storm + Peak scenario: approximately 20.39% energy-intensity reduction and 54.46% throughput gain in the most difficult tested condition.",
    ])

    add_heading(doc, "8. Challenges and Limitations", 1)
    add_bullets(doc, [
        "The environment is simulated and does not yet use real AIS, berth, crane, or grid-market data.",
        "The Q-learning policy is interpretable and lightweight but less expressive than a fully trained neural DQN.",
        "The reward function uses manually selected weights; real deployment would require stakeholder calibration.",
        "The Streamlit dashboard is suitable for demonstration but would need authentication, data pipelines, and live telemetry for production.",
    ])

    add_heading(doc, "9. Conclusion and Future Work", 1)
    add_body(doc, "The project successfully implements Port-Optimus as a complete reinforcement learning and MLOps demonstration. It includes a custom simulator, RL policy, manifest logging, results tracking, and a polished interactive dashboard. The evaluation shows consistent improvements in throughput, wait penalty, safety, and energy intensity compared with a fixed baseline.")
    add_bullets(doc, [
        "Future work 1: train the optional PyTorch DQN controller with replay memory and target networks.",
        "Future work 2: connect the simulator to real port arrival schedules and electricity market data.",
        "Future work 3: add crane scheduling, berth allocation, and multi-agent truck coordination.",
        "Future work 4: deploy the dashboard with authentication and experiment comparison history.",
    ])

    add_heading(doc, "10. References", 1)
    refs = [
        "Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction.",
        "Mnih, V., et al. (2015). Human-level control through deep reinforcement learning. Nature.",
        "United Nations. Sustainable Development Goals: SDG 9, SDG 11, and SDG 12.",
        "Streamlit Documentation. Building interactive data applications in Python.",
        "Plotly Documentation. Radar charts, heatmaps, and graph objects.",
        "Project source code: Port_Optimus_VS, local workspace implementation.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(ref)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
